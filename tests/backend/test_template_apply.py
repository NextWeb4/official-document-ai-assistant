import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET
from urllib.parse import unquote

import pytest
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

import config
from api.routes import templates as template_routes
from db.database import Base, get_db
from db.models import Document as DbDocument
from main import app
from core.document.font_utils import set_run_font
from core.document.parser import parse_docx
from core.document.template_applier import _verify_template_export, generate_docx_with_template


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": WORD_NS}


def _read_document_xml(path: Path) -> str:
    with zipfile.ZipFile(path, "r") as zf:
        return zf.read("word/document.xml").decode("utf-8")


def _make_content_doc(path: Path) -> Path:
    doc = Document()
    title = doc.add_paragraph()
    title.add_run("源文件标题")
    body = doc.add_paragraph()
    body.add_run("源文件正文123")
    doc.save(path)
    return path


def _make_template_doc(path: Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(21)
    section.bottom_margin = Mm(22)
    section.left_margin = Mm(23)
    section.right_margin = Mm(24)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("模板标题占位")
    set_run_font(title_run, "黑体", "宋体")
    title_run.font.size = Pt(20)

    body = doc.add_paragraph()
    body.paragraph_format.first_line_indent = Pt(28)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    body.paragraph_format.line_spacing = Pt(30)
    ppr = body._p.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:line"), "600")
    spacing.set(qn("w:lineRule"), "exact")
    body_run = body.add_run("模板正文占位")
    set_run_font(body_run, "楷体_GB2312", "Times New Roman")
    body_run.font.size = Pt(14)

    doc.save(path)
    return path


def _make_table_content_doc(path: Path, rows: int, cols: int) -> Path:
    doc = Document()
    doc.add_paragraph("源表格前正文")
    table = doc.add_table(rows=rows, cols=cols)
    for row in range(rows):
        for col in range(cols):
            table.cell(row, col).text = f"源单元格-{row}-{col}"
    doc.add_paragraph("源表格后正文")
    doc.save(path)
    return path


def _make_table_template_doc(path: Path, dimensions: list[tuple[int, int]]) -> Path:
    doc = Document()
    doc.add_paragraph("模板表格前占位")
    for table_index, (rows, cols) in enumerate(dimensions):
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in range(rows):
            for col in range(cols):
                table.cell(row, col).text = f"模板表{table_index}-占位-{row}-{col}"
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D9EAD3")
        table.cell(0, 0)._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph("模板表格后占位")
    doc.save(path)
    return path


def test_generate_docx_with_template_uses_template_format_and_content_text(tmp_path):
    content_path = _make_content_doc(tmp_path / "content.docx")
    template_path = _make_template_doc(tmp_path / "template.docx")
    output_path = tmp_path / "templated_output.docx"

    content_model = parse_docx(content_path)
    generate_docx_with_template(content_model, template_path, output_path)

    output = Document(str(output_path))
    text = "\n".join(p.text for p in output.paragraphs)
    assert "源文件标题" in text
    assert "源文件正文123" in text
    assert "模板标题占位" not in text
    assert "模板正文占位" not in text
    assert abs(output.sections[0].top_margin.mm - 21) < 1
    assert abs(output.sections[0].left_margin.mm - 23) < 1

    xml = _read_document_xml(output_path)
    root = ET.fromstring(xml)
    body_para = next(p for p in root.findall(".//w:p", NS) if p.find('.//w:t[.="源文件正文123"]', NS) is not None)
    rfonts = body_para.find(".//w:rFonts", NS)
    assert rfonts.get(f"{{{WORD_NS}}}eastAsia") == "楷体_GB2312"
    assert rfonts.get(f"{{{WORD_NS}}}ascii") == "Times New Roman"
    indent = body_para.find("./w:pPr/w:ind", NS)
    spacing = body_para.find("./w:pPr/w:spacing", NS)
    assert indent.get(f"{{{WORD_NS}}}firstLine") == "560"
    assert spacing.get(f"{{{WORD_NS}}}line") == "600"


def test_template_apply_rebuilds_expanded_table_and_removes_surplus_placeholders(tmp_path):
    content_path = _make_table_content_doc(tmp_path / "content_tables.docx", rows=2, cols=3)
    template_path = _make_table_template_doc(
        tmp_path / "template_tables.docx",
        dimensions=[(1, 1), (2, 2)],
    )
    output_path = tmp_path / "expanded_table_output.docx"

    content_model = parse_docx(content_path)
    generate_docx_with_template(content_model, template_path, output_path)

    output = Document(str(output_path))
    assert [paragraph.text for paragraph in output.paragraphs] == [
        "源表格前正文",
        "源表格后正文",
    ]
    assert len(output.tables) == 1
    table = output.tables[0]
    assert (len(table.rows), len(table.columns)) == (2, 3)
    assert [
        [table.cell(row, col).text for col in range(3)]
        for row in range(2)
    ] == [
        ["源单元格-0-0", "源单元格-0-1", "源单元格-0-2"],
        ["源单元格-1-0", "源单元格-1-1", "源单元格-1-2"],
    ]
    assert table.style.name == "Table Grid"
    assert table.alignment == WD_TABLE_ALIGNMENT.CENTER
    shading = table.cell(0, 0)._tc.tcPr.find(qn("w:shd"))
    assert shading is not None
    assert shading.get(qn("w:fill")) == "D9EAD3"
    assert "模板表" not in "\n".join(cell.text for row in table.rows for cell in row.cells)


def test_template_apply_shrinks_placeholder_table_to_source_dimensions(tmp_path):
    content_path = _make_table_content_doc(tmp_path / "small_content_table.docx", rows=1, cols=2)
    template_path = _make_table_template_doc(
        tmp_path / "large_template_table.docx",
        dimensions=[(3, 4)],
    )
    output_path = tmp_path / "shrunk_table_output.docx"

    content_model = parse_docx(content_path)
    generate_docx_with_template(content_model, template_path, output_path)

    output = Document(str(output_path))
    assert len(output.tables) == 1
    assert (len(output.tables[0].rows), len(output.tables[0].columns)) == (1, 2)
    assert [cell.text for cell in output.tables[0].rows[0].cells] == [
        "源单元格-0-0",
        "源单元格-0-1",
    ]


def test_template_export_verification_rejects_table_dimension_and_cell_loss(tmp_path):
    content_path = _make_table_content_doc(tmp_path / "verified_content.docx", rows=2, cols=2)
    template_path = _make_table_template_doc(tmp_path / "verified_template.docx", dimensions=[(1, 1)])
    content_model = parse_docx(content_path)
    template_model = parse_docx(template_path)

    wrong_dimensions = Document()
    wrong_dimensions.add_paragraph("源表格前正文")
    wrong_dimensions.add_table(rows=1, cols=1).cell(0, 0).text = "源单元格-0-0"
    wrong_dimensions.add_paragraph("源表格后正文")
    wrong_dimensions_path = tmp_path / "wrong_dimensions.docx"
    wrong_dimensions.save(wrong_dimensions_path)
    with pytest.raises(ValueError, match="表格维度复核失败"):
        _verify_template_export(wrong_dimensions_path, content_model, template_model)

    wrong_cell = Document()
    wrong_cell.add_paragraph("源表格前正文")
    table = wrong_cell.add_table(rows=2, cols=2)
    for row in range(2):
        for col in range(2):
            table.cell(row, col).text = f"源单元格-{row}-{col}"
    table.cell(1, 1).text = "被静默丢失"
    wrong_cell.add_paragraph("源表格后正文")
    wrong_cell_path = tmp_path / "wrong_cell.docx"
    wrong_cell.save(wrong_cell_path)
    with pytest.raises(ValueError, match=r"单元格 \(1,1\) 内容复核失败"):
        _verify_template_export(wrong_cell_path, content_model, template_model)


def test_uploaded_template_apply_download_api(tmp_path, monkeypatch):
    uploads_dir = tmp_path / "uploads"
    outputs_dir = tmp_path / "outputs"
    template_dir = tmp_path / "uploaded_templates"
    temp_dir = tmp_path / "temp"
    uploads_dir.mkdir()
    outputs_dir.mkdir()
    template_dir.mkdir()
    monkeypatch.setattr(config, "OUTPUT_DIR", outputs_dir)
    monkeypatch.setattr(config, "TEMP_DIR", temp_dir)
    monkeypatch.setattr(template_routes, "_UPLOADED_TEMPLATE_FILES_DIR", template_dir)
    monkeypatch.setattr(template_routes, "_UPLOADED_TEMPLATE_FILES_INDEX", template_dir / "index.json")

    engine = create_engine("sqlite:///:memory:", connect_args={"check_same_thread": False})
    Base.metadata.create_all(bind=engine)
    TestingSession = sessionmaker(autocommit=False, autoflush=False, bind=engine)
    db = TestingSession()

    content_path = _make_content_doc(uploads_dir / "content.docx")
    doc = DbDocument(
        filename="源文件.docx",
        file_path=str(content_path),
        file_hash="hash",
        document_type="notice",
        status="uploaded",
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    def override_get_db():
        try:
            yield db
        finally:
            pass

    app.dependency_overrides[get_db] = override_get_db
    client = TestClient(app)
    try:
        template_path = _make_template_doc(tmp_path / "福建模板.docx")
        with open(template_path, "rb") as fh:
            upload_resp = client.post(
                "/api/templates/files/upload",
                files={"file": ("福建模板.docx", fh, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
        assert upload_resp.status_code == 200
        template_id = upload_resp.json()["template"]["id"]
        assert template_id.startswith("tpl_")

        apply_resp = client.post(f"/api/templates/files/{template_id}/apply/{doc.id}/download")
        assert apply_resp.status_code == 200
        assert "源文件_套用模板_" in unquote(apply_resp.headers["content-disposition"])

        output = tmp_path / "api_templated.docx"
        output.write_bytes(apply_resp.content)
        xml = _read_document_xml(output)
        assert "源文件正文123" in xml
        assert "模板正文占位" not in xml

        generated = Document(str(output))
        assert abs(generated.sections[0].top_margin.mm - 21) < 1
        assert list(outputs_dir.iterdir()) == []
        assert list(temp_dir.iterdir()) == []

        preview_resp = client.post(
            f"/api/templates/files/{template_id}/apply-preview/download",
            json={
                "paragraphs": [
                    {"text": "当前预览修改后标题", "role": "title", "is_heading": True, "heading_level": 0},
                    {"text": "当前预览修改后正文456", "role": "body", "is_heading": False},
                ],
                "source_filename": "当前预览.docx",
            },
        )
        assert preview_resp.status_code == 200
        preview_output = tmp_path / "preview_templated.docx"
        preview_output.write_bytes(preview_resp.content)
        preview_xml = _read_document_xml(preview_output)
        assert "当前预览修改后正文456" in preview_xml
        assert "源文件正文123" not in preview_xml
        assert list(temp_dir.iterdir()) == []
    finally:
        app.dependency_overrides.pop(get_db, None)
        db.close()
