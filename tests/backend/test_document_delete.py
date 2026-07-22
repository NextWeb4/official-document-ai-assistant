from pathlib import Path
import re

import pytest

from docx import Document as WordDocument
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from db.database import Base
from db.models import Document
from services import document_service as svc


def _session():
    engine = create_engine("sqlite:///:memory:", connect_args={"check_same_thread": False})
    Base.metadata.create_all(bind=engine)
    return sessionmaker(autocommit=False, autoflush=False, bind=engine)()


def _make_docx(path: Path, text: str = "测试文档") -> Path:
    doc = WordDocument()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def test_upload_same_filename_uses_unique_storage_paths(tmp_path, monkeypatch):
    uploads = tmp_path / "uploads"
    outputs = tmp_path / "outputs"
    uploads.mkdir()
    outputs.mkdir()
    monkeypatch.setattr(svc, "UPLOAD_DIR", uploads)
    monkeypatch.setattr(svc, "OUTPUT_DIR", outputs)

    db = _session()
    first = _make_docx(tmp_path / "first.docx", "第一份")
    second = _make_docx(tmp_path / "second.docx", "第二份")

    doc1 = svc.upload_document(db, first, "同名.docx")
    doc2 = svc.upload_document(db, second, "同名.docx")

    assert doc1.file_path != doc2.file_path
    assert Path(doc1.file_path).exists()
    assert Path(doc2.file_path).exists()


def test_export_filename_uses_source_stem_and_unique_timestamp():
    first = svc.export_filename("原始文件.docx", "optimized")
    second = svc.export_filename("原始文件.docx", "optimized")

    pattern = r"原始文件_optimized_\d{8}_\d{6}_\d{6}_[0-9a-f]{8}\.docx"
    assert re.fullmatch(pattern, first)
    assert re.fullmatch(pattern, second)
    assert first != second


def test_delete_document_removes_record_and_owned_file(tmp_path, monkeypatch):
    uploads = tmp_path / "uploads"
    outputs = tmp_path / "outputs"
    uploads.mkdir()
    outputs.mkdir()
    monkeypatch.setattr(svc, "UPLOAD_DIR", uploads)
    monkeypatch.setattr(svc, "OUTPUT_DIR", outputs)

    db = _session()
    source = _make_docx(tmp_path / "source.docx")
    doc = svc.upload_document(db, source, "待删除.docx")
    stored_path = Path(doc.file_path)

    assert svc.delete_document(db, doc.id) is True
    assert db.query(Document).filter(Document.id == doc.id).first() is None
    assert not stored_path.exists()


def test_delete_document_keeps_shared_file_reference(tmp_path, monkeypatch):
    uploads = tmp_path / "uploads"
    outputs = tmp_path / "outputs"
    uploads.mkdir()
    outputs.mkdir()
    monkeypatch.setattr(svc, "UPLOAD_DIR", uploads)
    monkeypatch.setattr(svc, "OUTPUT_DIR", outputs)

    db = _session()
    shared = _make_docx(uploads / "shared.docx")
    doc1 = Document(filename="a.docx", file_path=str(shared), file_hash="a")
    doc2 = Document(filename="b.docx", file_path=str(shared), file_hash="b")
    db.add_all([doc1, doc2])
    db.commit()
    db.refresh(doc1)

    assert svc.delete_document(db, doc1.id) is True
    assert shared.exists()


def test_delete_document_file_failure_keeps_retryable_ownership(tmp_path, monkeypatch):
    uploads = tmp_path / "uploads"
    outputs = tmp_path / "outputs"
    uploads.mkdir()
    outputs.mkdir()
    monkeypatch.setattr(svc, "UPLOAD_DIR", uploads)
    monkeypatch.setattr(svc, "OUTPUT_DIR", outputs)

    db = _session()
    source = _make_docx(tmp_path / "source.docx")
    doc = svc.upload_document(db, source, "占用中.docx")
    doc_id = doc.id
    original_unlink = Path.unlink

    def fail_staged_unlink(self, *args, **kwargs):
        if self.name.endswith(".deleting"):
            raise PermissionError("file is in use")
        return original_unlink(self, *args, **kwargs)

    monkeypatch.setattr(Path, "unlink", fail_staged_unlink)

    with pytest.raises(svc.DocumentDeletionError, match="关闭占用"):
        svc.delete_document(db, doc_id)

    db.expire_all()
    preserved = db.query(Document).filter(Document.id == doc_id).one()
    preserved_path = Path(preserved.file_path)
    assert preserved_path.name.endswith(".deleting")
    assert preserved_path.exists()
