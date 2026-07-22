from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document as WordDocument
from fastapi import HTTPException
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

import config
from api.routes import ai as ai_routes
from api.routes import documents as document_routes
from api.routes import optimize as optimize_routes
from api.schemas.api_models import AISuggestion, ApplyAIRequest
from core.document.parser import parse_docx
from db.database import Base
from db.models import AIConfig, Document
from services import document_service as svc
from utils.crypto import encrypt_value


def _session():
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
    )
    Base.metadata.create_all(bind=engine)
    return sessionmaker(autocommit=False, autoflush=False, bind=engine)()


def _write_docx(path: Path, text: str) -> Path:
    document = WordDocument()
    document.add_paragraph(text)
    document.save(path)
    return path


def _docx_text(path: Path) -> str:
    return "\n".join(paragraph.text for paragraph in parse_docx(path).paragraphs)


def _add_document(db, source: Path, *, filename: str = "same.docx", optimized: Path | None = None) -> Document:
    document = Document(
        filename=filename,
        file_path=str(source),
        file_hash=f"hash-{source.name}-{id(source)}",
        document_type="notice",
        status="optimized" if optimized else "uploaded",
        optimized_path=str(optimized) if optimized else None,
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    return document


@pytest.mark.asyncio
async def test_sequential_docx_operations_compose_from_owned_current_output(tmp_path, monkeypatch):
    output_dir = tmp_path / "outputs"
    output_dir.mkdir()
    source = _write_docx(tmp_path / "source.docx", "base")
    db = _session()
    document = _add_document(db, source)
    monkeypatch.setattr(svc, "OUTPUT_DIR", output_dir)
    monkeypatch.setattr(config, "OUTPUT_DIR", output_dir)

    optimized_inputs: list[str] = []
    checked_inputs: list[str] = []

    class SequentialRuleEngine:
        def check_and_fix_with_count(self, model, _doc_type, _selected_rule_ids):
            paragraph = model.paragraphs[0]
            optimized_inputs.append(paragraph.text)
            updated = f"{paragraph.text}|rule-{len(optimized_inputs)}"
            paragraph.text = updated
            paragraph.runs[0].text = updated
            for run in paragraph.runs[1:]:
                run.text = ""
            return [], model, 1

        def check(self, model, _doc_type):
            checked_inputs.append(model.paragraphs[0].text)
            return []

    monkeypatch.setattr(svc, "_rule_engine", SequentialRuleEngine())

    first = svc.optimize_document(db, document.id)
    first_output = Path(first["output_path"])
    second = svc.optimize_document(db, document.id)
    second_output = Path(second["output_path"])

    assert optimized_inputs == ["base", "base|rule-1"]
    assert not first_output.exists()
    assert _docx_text(second_output) == "base|rule-1|rule-2"

    svc.check_document(db, document.id)
    assert checked_inputs == [
        "base|rule-1",
        "base|rule-1|rule-2",
        "base|rule-1|rule-2",
    ]

    monkeypatch.setenv("APP_MODE", "online")
    db.add(AIConfig(
        provider="deepseek",
        api_key_encrypted=encrypt_value("test-key"),
        base_url="https://example.test/v1",
        model="test-model",
        is_active=True,
    ))
    db.commit()
    analyzed_texts: list[str] = []

    class FakeProvider:
        async def analyze(self, document_text: str):
            analyzed_texts.append(document_text)
            return SimpleNamespace(issues=[], raw_response="ok")

    monkeypatch.setattr(ai_routes, "create_provider", lambda *_args: FakeProvider())
    analysis = await ai_routes.ai_analyze(document.id, provider="deepseek", db=db)

    assert analysis["success"] is True
    assert analyzed_texts == ["base|rule-1|rule-2"]

    applied = await ai_routes.apply_ai_suggestions(
        document.id,
        ApplyAIRequest(suggestions=[AISuggestion(original="rule-2", suggestion="ai-final")]),
        db=db,
    )

    db.refresh(document)
    final_output = Path(document.optimized_path)
    assert applied["success"] is True
    assert final_output != second_output
    assert not second_output.exists()
    assert _docx_text(final_output) == "base|rule-1|ai-final"
    assert _docx_text(source) == "base"


@pytest.mark.asyncio
async def test_ai_verification_exception_preserves_previous_committed_output(tmp_path, monkeypatch):
    output_dir = tmp_path / "outputs"
    output_dir.mkdir()
    source = _write_docx(tmp_path / "source.docx", "original")
    previous = _write_docx(output_dir / "previous.docx", "current text")
    db = _session()
    document = _add_document(db, source, optimized=previous)
    monkeypatch.setattr(svc, "OUTPUT_DIR", output_dir)
    monkeypatch.setattr(config, "OUTPUT_DIR", output_dir)
    monkeypatch.setattr(ai_routes, "export_filename", lambda *_args: "new-output.docx")

    from core.document import parser as parser_module

    real_parse = parser_module.parse_docx
    new_output = output_dir / "new-output.docx"

    def fail_new_output_verification(path):
        if Path(path) == new_output:
            raise RuntimeError("verification parser failed")
        return real_parse(path)

    monkeypatch.setattr(parser_module, "parse_docx", fail_new_output_verification)

    with pytest.raises(HTTPException) as exc_info:
        await ai_routes.apply_ai_suggestions(
            document.id,
            ApplyAIRequest(suggestions=[
                AISuggestion(original="current text", suggestion="verified text"),
            ]),
            db=db,
        )

    assert exc_info.value.status_code == 500
    assert "复核失败" in exc_info.value.detail
    db.expire_all()
    stored = db.query(Document).filter(Document.id == document.id).one()
    assert stored.optimized_path == str(previous)
    assert stored.status == "optimized"
    assert previous.exists()
    assert _docx_text(previous) == "current text"
    assert not new_output.exists()
    assert list(output_dir.iterdir()) == [previous]


@pytest.mark.asyncio
async def test_same_stem_documents_serve_only_their_explicit_owned_outputs(tmp_path, monkeypatch):
    output_dir = tmp_path / "outputs"
    output_dir.mkdir()
    first_source = _write_docx(tmp_path / "first-source.docx", "first source")
    second_source = _write_docx(tmp_path / "second-source.docx", "second source")
    first_output = _write_docx(output_dir / "same-first.docx", "first output")
    second_output = _write_docx(output_dir / "same-second.docx", "second output")
    db = _session()
    first = _add_document(db, first_source, optimized=first_output)
    second = _add_document(db, second_source, optimized=second_output)
    monkeypatch.setattr(svc, "OUTPUT_DIR", output_dir)

    first_download = await optimize_routes.download_optimized(first.id, db=db)
    second_download = await optimize_routes.download_optimized(second.id, db=db)
    first_preview = await document_routes.get_document_preview(first.id, db=db)
    second_preview = await document_routes.get_document_preview(second.id, db=db)

    assert Path(first_download.path) == first_output
    assert Path(second_download.path) == second_output
    assert first_preview["paragraphs"][0]["text"] == "first output"
    assert second_preview["paragraphs"][0]["text"] == "second output"


@pytest.mark.asyncio
async def test_same_stem_fallbacks_cannot_serve_another_documents_output(tmp_path, monkeypatch):
    output_dir = tmp_path / "outputs"
    output_dir.mkdir()
    first_source = _write_docx(tmp_path / "first-source.docx", "first source")
    second_source = _write_docx(tmp_path / "second-source.docx", "second source")
    other_output = _write_docx(output_dir / "same_optimized.docx", "second output")
    db = _session()
    no_output = _add_document(db, first_source)
    other = _add_document(db, second_source, optimized=other_output)
    stale = _add_document(
        db,
        first_source,
        optimized=tmp_path / "old-output-root" / other_output.name,
    )
    monkeypatch.setattr(svc, "OUTPUT_DIR", output_dir)

    for document in (no_output, stale):
        with pytest.raises(HTTPException) as download_error:
            await optimize_routes.download_optimized(document.id, db=db)
        assert download_error.value.status_code == 404

        with pytest.raises(HTTPException) as preview_error:
            await document_routes.get_document_preview(document.id, db=db)
        assert preview_error.value.status_code == 404

    assert other.optimized_path == str(other_output)
    assert other_output.exists()
