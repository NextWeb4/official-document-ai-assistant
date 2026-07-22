from pathlib import Path

from docx import Document

from api.routes.ai import _apply_suggestions_to_model
from api.schemas.api_models import AISuggestion
from core.document.generator import generate_docx
from core.document.parser import parse_docx


def test_ai_apply_replaces_numeric_suggestion_and_verifies_generated_docx(tmp_path):
    source = tmp_path / "numeric_source.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("年度编号：")
    paragraph.add_run("2024")
    paragraph.add_run("号")
    doc.save(source)

    model = parse_docx(source)
    result = _apply_suggestions_to_model(
        model,
        [AISuggestion(original=2024, suggestion=2025)],
    )

    assert result["failed"] == []
    assert result["applied"][0]["original"] == "2024"
    assert model.paragraphs[0].text == "年度编号：2025号"

    output = tmp_path / "numeric_output.docx"
    generate_docx(model, output)
    generated_model = parse_docx(output)
    generated_text = "\n".join(p.text for p in generated_model.paragraphs)
    assert "年度编号：2025号" in generated_text
    assert "年度编号：2024号" not in generated_text


def test_ai_apply_reports_unmatched_suggestion(tmp_path):
    source = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("正文内容")
    doc.save(source)

    model = parse_docx(source)
    result = _apply_suggestions_to_model(
        model,
        [AISuggestion(original="不存在的原文", suggestion="替换后")],
    )

    assert result["applied"] == []
    assert result["failed"][0]["reason"] == "未在当前文档中匹配到原文片段"


def test_ai_apply_matches_full_width_digits(tmp_path):
    source = tmp_path / "full_width_source.docx"
    doc = Document()
    doc.add_paragraph("编号：２０２４号")
    doc.save(source)

    model = parse_docx(source)
    result = _apply_suggestions_to_model(
        model,
        [AISuggestion(original="2024", suggestion="2025")],
    )

    assert result["failed"] == []
    assert model.paragraphs[0].text == "编号：2025号"
