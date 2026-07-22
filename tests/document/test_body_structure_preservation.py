"""Regression coverage for body block order and inherited formatting."""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from core.document.generator import generate_docx
from core.document.modifier import remove_extra_blank_lines, replace_paragraph_text
from core.document.parser import parse_docx


def _body_blocks(path: Path) -> list[tuple[str, str]]:
    doc = Document(str(path))
    blocks: list[tuple[str, str]] = []
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            blocks.append(("paragraph", DocxParagraph(child, doc).text))
        elif child.tag == qn("w:tbl"):
            table = DocxTable(child, doc)
            blocks.append(("table", table.cell(0, 0).text))
    return blocks


def _add_external_hyperlink(paragraph, text: str, address: str) -> None:
    relationship_id = paragraph.part.relate_to(address, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def test_removing_consecutive_blank_paragraphs_keeps_later_text_after_table(tmp_path):
    source = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("before")
    doc.add_paragraph("")
    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "table"
    doc.add_paragraph("after")
    doc.save(str(source))

    model = parse_docx(source)
    assert model.tables[0].insert_after_index == 2

    remove_extra_blank_lines(model)

    assert [paragraph.text for paragraph in model.paragraphs] == ["before", "", "after"]
    assert model.tables[0].insert_after_index == 1

    output = generate_docx(model, tmp_path / "output.docx")
    assert _body_blocks(output) == [
        ("paragraph", "before"),
        ("paragraph", ""),
        ("table", "table"),
        ("paragraph", "after"),
    ]


def test_table_breaks_blank_paragraph_sequence(tmp_path):
    source = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("before")
    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "table"
    doc.add_paragraph("")
    doc.add_paragraph("after")
    doc.save(str(source))

    model = parse_docx(source)
    remove_extra_blank_lines(model)

    assert [paragraph.text for paragraph in model.paragraphs] == ["before", "", "", "after"]
    output = generate_docx(model, tmp_path / "output.docx")
    assert _body_blocks(output) == [
        ("paragraph", "before"),
        ("paragraph", ""),
        ("table", "table"),
        ("paragraph", ""),
        ("paragraph", "after"),
    ]


def test_whole_paragraph_replacement_removes_stale_hyperlink_and_relationship(tmp_path):
    source = tmp_path / "source.docx"
    doc = Document()
    paragraph = doc.add_paragraph("prefix ")
    _add_external_hyperlink(paragraph, "linked text", "https://example.invalid/original")
    doc.save(str(source))

    model = parse_docx(source)
    replace_paragraph_text(model, 0, "replacement")

    output = generate_docx(model, tmp_path / "output.docx")
    result = Document(str(output))
    assert result.paragraphs[0].text == "replacement"
    assert result.paragraphs[0]._p.find(qn("w:hyperlink")) is None
    assert all(rel.reltype != RT.HYPERLINK for rel in result.part.rels.values())


def test_unmodified_hyperlink_only_paragraph_is_not_duplicated(tmp_path):
    source = tmp_path / "source.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    _add_external_hyperlink(paragraph, "linked text", "https://example.invalid/original")
    doc.save(str(source))

    model = parse_docx(source)
    assert model.paragraphs[0].runs == []

    output = generate_docx(model, tmp_path / "output.docx")
    result = Document(str(output))
    assert result.paragraphs[0].text == "linked text"
    assert len(result.paragraphs[0]._p.findall(qn("w:hyperlink"))) == 1
    assert sum(rel.reltype == RT.HYPERLINK for rel in result.part.rels.values()) == 1


def test_inherited_formatting_remains_implicit_after_round_trip(tmp_path):
    source = tmp_path / "source.docx"
    doc = Document()
    style = doc.styles.add_style("Inherited Body", WD_STYLE_TYPE.PARAGRAPH)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Pt(16)
    style.font.bold = True
    style.font.italic = True
    style.font.underline = True
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run("inherits from its paragraph style")
    doc.save(str(source))

    model = parse_docx(source)
    paragraph_model = model.paragraphs[0]
    assert paragraph_model.format.alignment is None
    assert paragraph_model.format.space_before_pt is None
    assert paragraph_model.format.space_after_pt is None
    assert paragraph_model.format.line_spacing_pt is None
    assert paragraph_model.format.first_line_indent_pt is None
    assert paragraph_model.runs[0].format.bold is None
    assert paragraph_model.runs[0].format.italic is None
    assert paragraph_model.runs[0].format.underline is None

    output = generate_docx(model, tmp_path / "output.docx")
    result = Document(str(output))
    output_paragraph = result.paragraphs[0]
    paragraph_properties = output_paragraph._p.pPr
    assert paragraph_properties.find(qn("w:jc")) is None
    assert paragraph_properties.find(qn("w:spacing")) is None
    assert paragraph_properties.find(qn("w:ind")) is None

    run_properties = output_paragraph.runs[0]._r.rPr
    assert run_properties.find(qn("w:b")) is None
    assert run_properties.find(qn("w:i")) is None
    assert run_properties.find(qn("w:u")) is None
