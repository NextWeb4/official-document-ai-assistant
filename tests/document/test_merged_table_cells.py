"""Regression coverage for canonical merged-table cells."""

from pathlib import Path

from docx import Document

from core.document.generator import generate_docx
from core.document.models import Run
from core.document.parser import parse_docx


def _set_anchor_text(cell, text: str) -> None:
    cell.text = text
    paragraph = cell.paragraphs[0]
    paragraph.text = text
    paragraph.runs = [Run(index=0, text=text)]


def _anchor(model, row: int, col: int):
    return next(
        cell
        for cell in model.tables[0].cells
        if cell.row == row and cell.col == col
    )


def _create_horizontal_merge(path: Path) -> Path:
    document = Document()
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "horizontal-original"
    table.cell(0, 2).text = "right"
    table.cell(1, 0).text = "bottom-left"
    table.cell(1, 1).text = "bottom-middle"
    table.cell(1, 2).text = "bottom-right"
    document.save(path)
    return path


def _create_vertical_merge(path: Path) -> Path:
    document = Document()
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).merge(table.cell(2, 0)).text = "vertical-original"
    table.cell(0, 1).text = "top-right"
    table.cell(1, 1).text = "middle-right"
    table.cell(2, 1).text = "bottom-right"
    document.save(path)
    return path


def test_horizontal_merge_has_one_anchor_and_keeps_anchor_edit(tmp_path):
    source = _create_horizontal_merge(tmp_path / "horizontal.docx")
    model = parse_docx(source)

    table_model = model.tables[0]
    assert (table_model.rows, table_model.cols) == (2, 3)
    assert len(table_model.cells) == 5
    anchor = _anchor(model, 0, 0)
    assert (anchor.row_span, anchor.col_span) == (1, 2)
    assert not any(cell.row == 0 and cell.col == 1 for cell in table_model.cells)

    _set_anchor_text(anchor, "horizontal-edited")
    output = generate_docx(model, tmp_path / "horizontal-output.docx")

    result = Document(output)
    result_table = result.tables[0]
    assert result_table.cell(0, 0)._tc is result_table.cell(0, 1)._tc
    assert result_table.cell(0, 0).text == "horizontal-edited"
    assert result_table.cell(0, 2).text == "right"

    reparsed = parse_docx(output)
    reparsed_anchor = _anchor(reparsed, 0, 0)
    assert (reparsed_anchor.row_span, reparsed_anchor.col_span) == (1, 2)
    assert reparsed_anchor.text == "horizontal-edited"


def test_vertical_merge_has_one_anchor_and_keeps_anchor_edit(tmp_path):
    source = _create_vertical_merge(tmp_path / "vertical.docx")
    model = parse_docx(source)

    table_model = model.tables[0]
    assert (table_model.rows, table_model.cols) == (3, 2)
    assert len(table_model.cells) == 4
    anchor = _anchor(model, 0, 0)
    assert (anchor.row_span, anchor.col_span) == (3, 1)
    assert not any(
        cell.col == 0 and cell.row in {1, 2}
        for cell in table_model.cells
    )

    _set_anchor_text(anchor, "vertical-edited")
    output = generate_docx(model, tmp_path / "vertical-output.docx")

    result = Document(output)
    result_table = result.tables[0]
    assert result_table.cell(0, 0)._tc is result_table.cell(1, 0)._tc
    assert result_table.cell(0, 0)._tc is result_table.cell(2, 0)._tc
    assert result_table.cell(0, 0).text == "vertical-edited"
    assert result_table.cell(1, 1).text == "middle-right"

    reparsed = parse_docx(output)
    reparsed_anchor = _anchor(reparsed, 0, 0)
    assert (reparsed_anchor.row_span, reparsed_anchor.col_span) == (3, 1)
    assert reparsed_anchor.text == "vertical-edited"


def test_merge_geometry_is_rebuilt_without_the_source_file(tmp_path):
    source = _create_horizontal_merge(tmp_path / "source.docx")
    model = parse_docx(source)
    _set_anchor_text(_anchor(model, 0, 0), "rebuilt-anchor")
    model.source_path = None

    output = generate_docx(model, tmp_path / "rebuilt.docx")

    result = Document(output)
    table = result.tables[0]
    assert table.cell(0, 0)._tc is table.cell(0, 1)._tc
    assert table.cell(0, 0).text == "rebuilt-anchor"
    assert table.cell(0, 2).text == "right"


def test_vertical_merge_geometry_is_rebuilt_without_the_source_file(tmp_path):
    source = _create_vertical_merge(tmp_path / "source.docx")
    model = parse_docx(source)
    _set_anchor_text(_anchor(model, 0, 0), "rebuilt-vertical-anchor")
    model.source_path = None

    output = generate_docx(model, tmp_path / "rebuilt.docx")

    result = Document(output)
    table = result.tables[0]
    assert table.cell(0, 0)._tc is table.cell(1, 0)._tc
    assert table.cell(0, 0)._tc is table.cell(2, 0)._tc
    assert table.cell(0, 0).text == "rebuilt-vertical-anchor"
    assert table.cell(1, 1).text == "middle-right"
