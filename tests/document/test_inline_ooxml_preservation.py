"""Regression coverage for opaque inline OOXML during parse/generate round trips."""

from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from lxml import etree

from core.document.generator import generate_docx
from core.document.parser import parse_docx


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
DRAWING_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
NS = {"w": WORD_NS, "r": REL_NS, "a": DRAWING_NS}

_ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/"
    "x8AAusB9Wl5pWQAAAAASUVORK5CYII="
)


def _add_complex_field(paragraph, instruction: str, result: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._element.append(begin)

    instruction_run = paragraph.add_run()
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = f" {instruction} "
    instruction_run._element.append(instruction_node)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._element.append(separate)

    paragraph.add_run(result)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._element.append(end)


def _populate_story_paragraph(paragraph, label: str, instruction: str) -> None:
    paragraph.add_run(f"{label}-before")
    paragraph.add_run().add_picture(BytesIO(_ONE_PIXEL_PNG), width=Inches(0.1))
    paragraph.add_run(f"{label}-anchor")
    _add_complex_field(paragraph, instruction, f"{label}-result")
    paragraph.add_run(f"{label}-after")


def _create_inline_fixture(path: Path) -> Path:
    doc = Document()
    _populate_story_paragraph(doc.add_paragraph(), "body", 'DATE \\@ "yyyy-MM-dd"')

    table = doc.add_table(rows=1, cols=1)
    _populate_story_paragraph(table.cell(0, 0).paragraphs[0], "table", "REF bookmark_name")

    header = doc.sections[0].header.paragraphs[0]
    _populate_story_paragraph(header, "header", "SEQ Figure")

    footer = doc.sections[0].footer.paragraphs[0]
    _populate_story_paragraph(footer, "footer", 'DATE \\@ "yyyy"')
    footer.add_run("footer-page-anchor")
    _add_complex_field(footer, "PAGE", "1")

    doc.save(path)
    return path


def _read_xml(docx_path: Path, part_name: str):
    with ZipFile(docx_path) as package:
        return etree.fromstring(package.read(part_name))


def _find_paragraph(root, marker: str):
    for paragraph in root.xpath(".//w:p", namespaces=NS):
        text = "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))
        if marker in text:
            return paragraph
    raise AssertionError(f"Paragraph containing {marker!r} not found")


def _run_events(paragraph) -> list[tuple[str, str]]:
    events: list[tuple[str, str]] = []
    for run in paragraph.findall("w:r", NS):
        embeds = run.xpath(".//a:blip/@r:embed", namespaces=NS)
        field_chars = run.findall(".//w:fldChar", NS)
        instructions = run.findall(".//w:instrText", NS)
        if embeds:
            events.append(("drawing", embeds[0]))
        elif field_chars:
            events.append(("fldChar", field_chars[0].get(qn("w:fldCharType"))))
        elif instructions:
            events.append(("instrText", instructions[0].text or ""))
        else:
            events.append(("text", "".join(run.xpath(".//w:t/text()", namespaces=NS))))
    return events


def _atomic_events(paragraph) -> list[tuple[str, str]]:
    return [event for event in _run_events(paragraph) if event[0] != "text"]


def _relationship_ids(docx_path: Path, rels_part_name: str) -> set[str]:
    root = _read_xml(docx_path, rels_part_name)
    return {
        relationship.get("Id")
        for relationship in root.findall(f"{{{PACKAGE_REL_NS}}}Relationship")
    }


def _story_paragraphs(docx_path: Path) -> dict[str, object]:
    document_root = _read_xml(docx_path, "word/document.xml")
    header_root = _read_xml(docx_path, "word/header1.xml")
    footer_root = _read_xml(docx_path, "word/footer1.xml")
    return {
        "body": _find_paragraph(document_root, "body-anchor"),
        "table": _find_paragraph(document_root, "table-anchor"),
        "header": _find_paragraph(header_root, "header-anchor"),
        "footer": _find_paragraph(footer_root, "footer-anchor"),
    }


def _assert_embeds_resolve(docx_path: Path, paragraphs: dict[str, object]) -> None:
    body_rel_ids = _relationship_ids(docx_path, "word/_rels/document.xml.rels")
    header_rel_ids = _relationship_ids(docx_path, "word/_rels/header1.xml.rels")
    footer_rel_ids = _relationship_ids(docx_path, "word/_rels/footer1.xml.rels")

    for story in ("body", "table"):
        embed = next(value for kind, value in _run_events(paragraphs[story]) if kind == "drawing")
        assert embed in body_rel_ids
    header_embed = next(value for kind, value in _run_events(paragraphs["header"]) if kind == "drawing")
    footer_embed = next(value for kind, value in _run_events(paragraphs["footer"]) if kind == "drawing")
    assert header_embed in header_rel_ids
    assert footer_embed in footer_rel_ids


def test_format_only_roundtrip_preserves_inline_ooxml_in_all_story_parts(tmp_path):
    source = _create_inline_fixture(tmp_path / "inline-source.docx")
    output = tmp_path / "inline-output.docx"

    generate_docx(parse_docx(source), output)

    source_paragraphs = _story_paragraphs(source)
    output_paragraphs = _story_paragraphs(output)
    for story in ("body", "table", "header", "footer"):
        assert _run_events(output_paragraphs[story]) == _run_events(source_paragraphs[story])

    assert "DATE" in "".join(value for kind, value in _atomic_events(output_paragraphs["body"]) if kind == "instrText")
    assert "REF" in "".join(value for kind, value in _atomic_events(output_paragraphs["table"]) if kind == "instrText")
    assert "SEQ" in "".join(value for kind, value in _atomic_events(output_paragraphs["header"]) if kind == "instrText")
    footer_instructions = "".join(
        value for kind, value in _atomic_events(output_paragraphs["footer"]) if kind == "instrText"
    )
    assert "DATE" in footer_instructions
    assert "PAGE" in footer_instructions
    _assert_embeds_resolve(output, output_paragraphs)


def test_modeled_text_and_format_changes_preserve_atomic_inline_ooxml(tmp_path):
    source = _create_inline_fixture(tmp_path / "inline-source.docx")
    model = parse_docx(source)
    paragraphs = {
        "body": model.paragraphs[0],
        "table": model.tables[0].cells[0].paragraphs[0],
        "header": model.headers[0].paragraphs[0],
        "footer": model.footers[0].paragraphs[0],
    }
    for story, paragraph in paragraphs.items():
        run = next(run for run in paragraph.runs if run.text == f"{story}-before")
        run.text = f"{story}-updated"
        run.format.bold = True
        paragraph.text = paragraph.text.replace(f"{story}-before", f"{story}-updated")

    output = tmp_path / "inline-updated.docx"
    generate_docx(model, output)

    source_paragraphs = _story_paragraphs(source)
    output_paragraphs = _story_paragraphs(output)
    for story in ("body", "table", "header", "footer"):
        output_paragraph = output_paragraphs[story]
        assert _atomic_events(output_paragraph) == _atomic_events(source_paragraphs[story])
        updated_run = next(
            run
            for run in output_paragraph.findall("w:r", NS)
            if "".join(run.xpath(".//w:t/text()", namespaces=NS)) == f"{story}-updated"
        )
        assert updated_run.find("w:rPr/w:b", NS) is not None

        events = _run_events(output_paragraph)
        updated_idx = events.index(("text", f"{story}-updated"))
        drawing_idx = next(i for i, event in enumerate(events) if event[0] == "drawing")
        anchor_idx = events.index(("text", f"{story}-anchor"))
        instruction_idx = next(i for i, event in enumerate(events) if event[0] == "instrText")
        assert updated_idx < drawing_idx < anchor_idx < instruction_idx

    _assert_embeds_resolve(output, output_paragraphs)
