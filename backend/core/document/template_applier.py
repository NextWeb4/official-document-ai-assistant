# This file is part of the Official Document AI Assistant.
# (c) 2026 Jose AI (https://www.linhut.cn)
# Licensed under the MIT License. See the LICENSE file for details.
"""
Apply an uploaded Word template's formatting to another document's content.

Invariant:
- Text/tables come from the content document.
- Page setup, styles, headers/footers and paragraph/run formatting come from the template.
- The uploaded template file is read-only and never overwritten.
"""
from __future__ import annotations

import copy
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from core.document.generator import generate_docx
from core.document.models import (
    DocumentModel,
    Paragraph,
    ParagraphFormat,
    Run,
    RunFormat,
    Table,
    TableCell,
)
from core.document.parser import parse_docx


def normalize_word_template(template_path: Path, work_dir: Path | None = None) -> Path:
    """Return a .docx-compatible copy of a .docx/.dotx template."""
    template_path = Path(template_path)
    suffix = template_path.suffix.lower()
    if suffix == ".docx":
        return template_path
    if suffix != ".dotx":
        raise ValueError("Only .docx and .dotx templates are supported")

    target_dir = work_dir or Path(tempfile.mkdtemp(prefix="template_docx_"))
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / f"{template_path.stem}.docx"

    with zipfile.ZipFile(template_path, "r") as zin:
        with zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    text = data.decode("utf-8")
                    text = text.replace(
                        "wordprocessingml.template.main+xml",
                        "wordprocessingml.document.main+xml",
                    )
                    data = text.encode("utf-8")
                zout.writestr(item, data)
    return target


def generate_docx_with_template(
    content_model: DocumentModel,
    template_path: Path | str,
    output_path: Path | str,
) -> Path:
    """Generate a new .docx whose content comes from content_model and format from template_path."""
    template_path = Path(template_path)
    output_path = Path(output_path)
    with tempfile.TemporaryDirectory(prefix="template_apply_") as tmp:
        normalized_template = normalize_word_template(template_path, Path(tmp))
        template_model = parse_docx(str(normalized_template))
        prepared_template = Path(tmp) / "prepared_template.docx"
        _prepare_template_source(normalized_template, prepared_template, content_model)
        merged = apply_template_format(content_model, template_model, prepared_template)
        result = generate_docx(merged, output_path)
        _verify_template_export(result, content_model, template_model)
        return result


def apply_template_format(
    content_model: DocumentModel,
    template_model: DocumentModel,
    template_source_path: Path | str,
) -> DocumentModel:
    """Create a model using content text with template formatting."""
    merged = copy.deepcopy(content_model)
    merged.source_path = str(template_source_path)
    merged.page_setup = copy.deepcopy(template_model.page_setup)
    merged.headers = copy.deepcopy(template_model.headers)
    merged.footers = copy.deepcopy(template_model.footers)

    style_map = _build_style_map(template_model)
    fallback = style_map.get("body") or style_map.get("default")

    merged.paragraphs = [
        _paragraph_with_template_format(para, _template_for_para(para, style_map, fallback), idx)
        for idx, para in enumerate(content_model.paragraphs)
    ]
    merged.tables = [
        _table_with_template_format(
            table,
            _template_table_for_index(template_model, idx),
            style_map,
            fallback,
        )
        for idx, table in enumerate(content_model.tables)
    ]
    return merged


def _prepare_template_source(
    template_path: Path,
    prepared_path: Path,
    content_model: DocumentModel,
) -> None:
    """Create exact-size table shells while retaining the template's table formatting."""
    _validate_source_tables(content_model.tables)

    doc = Document(str(template_path))
    body = doc.element.body
    prototypes = list(doc.tables)
    template_table_elements = list(body.iter(qn("w:tbl")))

    # Table positions are paragraph-relative. Add blank shells when source content
    # has more paragraphs than the template so every source position is addressable.
    while len(doc.paragraphs) < len(content_model.paragraphs):
        doc.add_paragraph()
    paragraph_elements = [paragraph._p for paragraph in doc.paragraphs]

    prepared_tables = []
    for index, table_model in enumerate(content_model.tables):
        table = doc.add_table(rows=table_model.rows, cols=table_model.cols)
        if prototypes:
            prototype = prototypes[min(index, len(prototypes) - 1)]
            _copy_table_format(prototype, table)
        prepared_tables.append((table._tbl, table_model.insert_after_index))

    for table_element in reversed(template_table_elements):
        parent = table_element.getparent()
        if parent is not None:
            parent.remove(table_element)

    position_tails = {}
    for table_element, requested_position in prepared_tables:
        parent = table_element.getparent()
        if parent is not None:
            parent.remove(table_element)
        position = _normalized_table_position(requested_position, len(paragraph_elements))
        previous_at_position = position_tails.get(position)
        if previous_at_position is not None:
            previous_at_position.addnext(table_element)
        elif position < 0:
            if paragraph_elements:
                paragraph_elements[0].addprevious(table_element)
            else:
                body.insert(max(0, len(body) - 1), table_element)
        else:
            paragraph_elements[position].addnext(table_element)
        position_tails[position] = table_element

    doc.save(str(prepared_path))


def _validate_source_tables(tables: list[Table]) -> None:
    for table_index, table in enumerate(tables):
        if table.rows < 1 or table.cols < 1:
            raise ValueError(
                f"源文档第 {table_index + 1} 个表格维度无效：{table.rows}x{table.cols}"
            )

        coordinates: set[tuple[int, int]] = set()
        for cell in table.cells:
            coordinate = (cell.row, cell.col)
            if not (0 <= cell.row < table.rows and 0 <= cell.col < table.cols):
                raise ValueError(
                    "源文档表格单元格超出声明维度："
                    f"表格 {table_index + 1}，单元格 ({cell.row},{cell.col})，"
                    f"维度 {table.rows}x{table.cols}"
                )
            if coordinate in coordinates:
                raise ValueError(
                    "源文档表格包含重复单元格坐标："
                    f"表格 {table_index + 1}，单元格 ({cell.row},{cell.col})"
                )
            coordinates.add(coordinate)


def _copy_table_format(source, target) -> None:
    source_properties = source._tbl.tblPr
    target_properties = target._tbl.tblPr
    target_properties.getparent().replace(target_properties, copy.deepcopy(source_properties))

    source_grid_columns = source._tbl.tblGrid.findall(qn("w:gridCol"))
    target_grid_columns = target._tbl.tblGrid.findall(qn("w:gridCol"))
    if source_grid_columns:
        for index, target_column in enumerate(target_grid_columns):
            source_column = source_grid_columns[min(index, len(source_grid_columns) - 1)]
            width = source_column.get(qn("w:w"))
            if width is None:
                target_column.attrib.pop(qn("w:w"), None)
            else:
                target_column.set(qn("w:w"), width)

    source_rows = list(source.rows)
    if not source_rows:
        return

    for row_index, target_row in enumerate(target.rows):
        source_row = source_rows[min(row_index, len(source_rows) - 1)]
        _copy_row_format(source_row, target_row)
        source_cells = list(source_row.cells)
        if not source_cells:
            continue
        for col_index, target_cell in enumerate(target_row.cells):
            source_cell = source_cells[min(col_index, len(source_cells) - 1)]
            _copy_cell_format(source_cell, target_cell)


def _copy_row_format(source_row, target_row) -> None:
    source_properties = source_row._tr.trPr
    if source_properties is None:
        return
    properties = copy.deepcopy(source_properties)
    for tag in ("w:gridBefore", "w:gridAfter"):
        for element in properties.findall(qn(tag)):
            properties.remove(element)
    target_properties = target_row._tr.trPr
    if target_properties is None:
        target_row._tr.insert(0, properties)
    else:
        target_row._tr.replace(target_properties, properties)


def _copy_cell_format(source_cell, target_cell) -> None:
    properties = copy.deepcopy(source_cell._tc.tcPr)
    for tag in ("w:gridSpan", "w:hMerge", "w:vMerge"):
        for element in properties.findall(qn(tag)):
            properties.remove(element)
    target_properties = target_cell._tc.tcPr
    target_properties.getparent().replace(target_properties, properties)

    if not source_cell.paragraphs or not target_cell.paragraphs:
        return
    source_ppr = source_cell.paragraphs[0]._p.pPr
    if source_ppr is None:
        return
    target_paragraph = target_cell.paragraphs[0]._p
    target_ppr = target_paragraph.pPr
    if target_ppr is None:
        target_paragraph.insert(0, copy.deepcopy(source_ppr))
    else:
        target_paragraph.replace(target_ppr, copy.deepcopy(source_ppr))


def _normalized_table_position(position: int, paragraph_count: int) -> int:
    if paragraph_count == 0 or position < 0:
        return -1
    return min(position, paragraph_count - 1)


def _build_style_map(template_model: DocumentModel) -> dict[str, Paragraph]:
    style_map: dict[str, Paragraph] = {}
    for para in template_model.paragraphs:
        if not (para.text or "").strip() and not para.runs:
            continue
        key = _style_key(para)
        style_map.setdefault(key, para)
        style_map.setdefault("default", para)
    if "body" not in style_map:
        for para in template_model.paragraphs:
            if (para.text or "").strip() and not para.is_heading:
                style_map["body"] = para
                break
    return style_map


def _style_key(para: Paragraph) -> str:
    if para.role:
        return para.role
    if para.is_heading and para.heading_level is not None:
        return "title" if para.heading_level == 0 else f"heading_{para.heading_level}"
    return "body"


def _template_for_para(
    para: Paragraph,
    style_map: dict[str, Paragraph],
    fallback: Paragraph | None,
) -> Paragraph | None:
    key = _style_key(para)
    if key in style_map:
        return style_map[key]
    if para.is_heading and para.heading_level is not None:
        return style_map.get(f"heading_{para.heading_level}") or style_map.get("body") or fallback
    return style_map.get(para.role or "") or style_map.get("body") or fallback


def _paragraph_with_template_format(
    content_para: Paragraph,
    template_para: Paragraph | None,
    index: int,
) -> Paragraph:
    para = copy.deepcopy(content_para)
    para.index = index
    if template_para is None:
        return para

    para.format = copy.deepcopy(template_para.format or ParagraphFormat())
    template_run_format = _first_run_format(template_para)
    para.runs = [Run(index=0, text=content_para.text or "", format=copy.deepcopy(template_run_format))]
    return para


def _first_run_format(para: Paragraph) -> RunFormat:
    for run in para.runs:
        if (run.text or "").strip():
            return copy.deepcopy(run.format)
    if para.runs:
        return copy.deepcopy(para.runs[0].format)
    return RunFormat()


def _table_with_template_format(
    table: Table,
    template_table: Table | None,
    style_map: dict[str, Paragraph],
    fallback: Paragraph | None,
) -> Table:
    templated = copy.deepcopy(table)
    body_style = style_map.get("body") or fallback
    header_style = style_map.get("heading_1") or body_style
    template_cells = (
        {(cell.row, cell.col): cell for cell in template_table.cells}
        if template_table is not None
        else {}
    )
    for cell in templated.cells:
        template_cell = _template_cell_for_position(cell.row, cell.col, template_table, template_cells)
        cell_styles = template_cell.paragraphs if template_cell is not None else []
        fallback_style = header_style if cell.row == 0 else body_style
        cell.paragraphs = [
            _paragraph_with_template_format(
                para,
                cell_styles[min(idx, len(cell_styles) - 1)] if cell_styles else fallback_style,
                idx,
            )
            for idx, para in enumerate(cell.paragraphs)
        ]
    return templated


def _template_table_for_index(template_model: DocumentModel, index: int) -> Table | None:
    if not template_model.tables:
        return None
    return template_model.tables[min(index, len(template_model.tables) - 1)]


def _template_cell_for_position(
    row: int,
    col: int,
    template_table: Table | None,
    cells: dict[tuple[int, int], TableCell],
) -> TableCell | None:
    if template_table is None or template_table.rows < 1 or template_table.cols < 1:
        return None
    position = (min(row, template_table.rows - 1), min(col, template_table.cols - 1))
    return cells.get(position)


def _verify_template_export(
    output_path: Path,
    content_model: DocumentModel,
    template_model: DocumentModel,
) -> None:
    doc = Document(str(output_path))
    expected_paragraphs = [paragraph.text or "" for paragraph in content_model.paragraphs]
    actual_paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    if actual_paragraphs != expected_paragraphs:
        difference = _first_sequence_difference(expected_paragraphs, actual_paragraphs)
        raise ValueError(f"模板套用后正文段落复核失败：{difference}")

    body_table_count = sum(1 for _ in doc.element.body.iter(qn("w:tbl")))
    if body_table_count != len(content_model.tables) or len(doc.tables) != len(content_model.tables):
        raise ValueError(
            "模板套用后表格数量复核失败："
            f"实际 {body_table_count}，期望 {len(content_model.tables)}"
        )

    for table_index, (actual_table, expected_table) in enumerate(
        zip(doc.tables, content_model.tables),
        start=1,
    ):
        actual_dimensions = (len(actual_table.rows), len(actual_table.columns))
        expected_dimensions = (expected_table.rows, expected_table.cols)
        if actual_dimensions != expected_dimensions:
            raise ValueError(
                f"模板套用后第 {table_index} 个表格维度复核失败："
                f"实际 {actual_dimensions[0]}x{actual_dimensions[1]}，"
                f"期望 {expected_dimensions[0]}x{expected_dimensions[1]}"
            )

        expected_cells = {
            (cell.row, cell.col): _expected_cell_text(cell)
            for cell in expected_table.cells
        }
        for row in range(expected_table.rows):
            for col in range(expected_table.cols):
                expected_text = expected_cells.get((row, col), "")
                actual_text = actual_table.cell(row, col).text
                if actual_text != expected_text:
                    raise ValueError(
                        f"模板套用后第 {table_index} 个表格单元格 ({row},{col}) "
                        f"内容复核失败：实际 {actual_text!r}，期望 {expected_text!r}"
                    )

    section = doc.sections[0]
    expected = template_model.page_setup
    checks = [
        (expected.margin_top_mm, section.top_margin.mm, "上边距"),
        (expected.margin_bottom_mm, section.bottom_margin.mm, "下边距"),
        (expected.margin_left_mm, section.left_margin.mm, "左边距"),
        (expected.margin_right_mm, section.right_margin.mm, "右边距"),
    ]
    for expected_mm, actual_mm, label in checks:
        if expected_mm is not None and abs(actual_mm - expected_mm) > 1.0:
            raise ValueError(f"模板套用后{label}复核失败：{actual_mm:.1f}mm，期望 {expected_mm:.1f}mm")


def _expected_cell_text(cell) -> str:
    if cell.paragraphs:
        return "\n".join(paragraph.text or "" for paragraph in cell.paragraphs)
    return cell.text or ""


def _first_sequence_difference(expected: list[str], actual: list[str]) -> str:
    for index, (expected_text, actual_text) in enumerate(zip(expected, actual), start=1):
        if expected_text != actual_text:
            return f"第 {index} 段实际 {actual_text!r}，期望 {expected_text!r}"
    return f"实际 {len(actual)} 段，期望 {len(expected)} 段"


def copy_template_to_docx(src_path: Path, dst_path: Path) -> None:
    """Copy a .docx/.dotx template into a persistent .docx file."""
    src_path = Path(src_path)
    dst_path = Path(dst_path)
    dst_path.parent.mkdir(parents=True, exist_ok=True)
    if src_path.suffix.lower() == ".docx":
        shutil.copy2(src_path, dst_path)
        return
    with tempfile.TemporaryDirectory(prefix="template_upload_") as tmp:
        normalized = normalize_word_template(src_path, Path(tmp))
        shutil.copy2(normalized, dst_path)
